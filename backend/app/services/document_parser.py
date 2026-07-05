from pathlib import Path
import fitz
from docx import Document as DocxDocument
import pandas as pd


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".csv",
    ".xlsx",
    ".xls",
}


def clean_text(text: str) -> str:
    if not text:
        return ""

    return " ".join(text.replace("\x00", " ").split())


def parse_txt(file_path: str) -> str:
    path = Path(file_path)

    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return ""


def parse_pdf(file_path: str) -> str:
    text_parts = []

    try:
        pdf = fitz.open(file_path)

        for page in pdf:
            text_parts.append(page.get_text())

        pdf.close()
    except Exception:
        return ""

    return "\n".join(text_parts)


def parse_docx(file_path: str) -> str:
    text_parts = []

    try:
        document = DocxDocument(file_path)

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
    except Exception:
        return ""

    return "\n".join(text_parts)


def dataframe_to_text(df: pd.DataFrame, max_rows: int = 40) -> str:
    if df is None or df.empty:
        return ""

    df = df.fillna("")
    rows = []

    columns = [str(column).strip() for column in df.columns]
    rows.append("Columns: " + " | ".join(columns))

    for index, row in df.head(max_rows).iterrows():
        values = [str(value).strip() for value in row.tolist()]
        rows.append("Row " + str(index + 1) + ": " + " | ".join(values))

    rows.append(f"Table summary: {len(df)} rows, {len(df.columns)} columns.")

    return "\n".join(rows)


def parse_csv(file_path: str) -> str:
    try:
        df = pd.read_csv(file_path)
        return dataframe_to_text(df)
    except Exception as error:
        return f"Table extraction failed: {str(error)}"


def parse_excel(file_path: str) -> str:
    try:
        excel = pd.ExcelFile(file_path)
        text_parts = []

        for sheet_name in excel.sheet_names[:3]:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            sheet_text = dataframe_to_text(df)

            if sheet_text:
                text_parts.append(f"Sheet: {sheet_name}\n{sheet_text}")

        return "\n\n".join(text_parts)
    except Exception as error:
        return f"Table extraction failed: {str(error)}"


def extract_text_from_file(file_path: str) -> str:
    extension = Path(file_path).suffix.lower()

    if extension == ".txt":
        return clean_text(parse_txt(file_path))

    if extension == ".pdf":
        return clean_text(parse_pdf(file_path))

    if extension == ".docx":
        return clean_text(parse_docx(file_path))

    if extension == ".csv":
        return clean_text(parse_csv(file_path))

    if extension in [".xlsx", ".xls"]:
        return clean_text(parse_excel(file_path))

    return ""


def build_text_preview(text: str, max_length: int = 500) -> str:
    if not text:
        return "No readable text extracted."

    if len(text) <= max_length:
        return text

    return text[:max_length].strip() + "..."
