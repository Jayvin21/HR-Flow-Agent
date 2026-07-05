from pathlib import Path
import csv

BASE_DIR = Path(__file__).parent

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from openpyxl import Workbook
except ImportError:
    Workbook = None

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
except ImportError:
    canvas = None
    letter = None


def write_txt(path: Path, content: str):
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content.strip(), encoding="utf-8")


def write_docx(path: Path, title: str, paragraphs: list[str]):
    if Document is None:
        print(f"Skipping DOCX, python-docx not installed: {path}")
        return

    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading(title, level=1)

    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)

    doc.save(path)


def write_pdf(path: Path, title: str, lines: list[str]):
    if canvas is None:
        print(f"Skipping PDF, reportlab not installed: {path}")
        return

    path.parent.mkdir(parents=True, exist_ok=True)

    pdf = canvas.Canvas(str(path), pagesize=letter)
    width, height = letter

    y = height - 72
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(72, y, title)

    y -= 32
    pdf.setFont("Helvetica", 10)

    for line in lines:
        if y < 72:
            pdf.showPage()
            pdf.setFont("Helvetica", 10)
            y = height - 72

        pdf.drawString(72, y, line[:95])
        y -= 16

    pdf.save()


def write_csv(path: Path, headers: list[str], rows: list[list[str]]):
    path.parent.mkdir(parents=True, exist_ok=True)

    with path.open("w", newline="", encoding="utf-8") as file:
        writer = csv.writer(file)
        writer.writerow(headers)
        writer.writerows(rows)


def write_xlsx(path: Path, headers: list[str], rows: list[list[str]]):
    if Workbook is None:
        print(f"Skipping XLSX, openpyxl not installed: {path}")
        return

    path.parent.mkdir(parents=True, exist_ok=True)

    wb = Workbook()
    ws = wb.active
    ws.title = "Attendance"

    ws.append(headers)

    for row in rows:
        ws.append(row)

    wb.save(path)


# -----------------------------
# TXT FILES
# -----------------------------

write_txt(
    BASE_DIR / "resumes" / "arjun_mehta_resume.txt",
    """
Arjun Mehta
Email: arjun.mehta@email.com
Phone: +91 9876543210
Role: Frontend Developer

Summary:
Frontend developer with 2 years of experience building React dashboards,
admin panels, and SaaS interfaces.

Skills:
React, Next.js, JavaScript, TypeScript, Tailwind CSS, REST APIs, Git, MongoDB.

Projects:
1. HR Dashboard - Built employee analytics dashboard using React and Tailwind.
2. E-commerce Admin Panel - Implemented product, order, and inventory modules.

Experience:
Frontend Developer Intern at WebCraft Labs.
Worked on reusable components, API integration, form validation, and dashboard UI.

Education:
B.E. Computer Engineering.
"""
)

write_txt(
    BASE_DIR / "job_descriptions" / "frontend_developer_jd.txt",
    """
Frontend Developer Job Description

Required Skills:
React, Next.js, JavaScript, Tailwind CSS, REST API integration, Git.

Preferred Skills:
TypeScript, dashboard UI, charting libraries, authentication, SaaS product experience.

Responsibilities:
Build clean and responsive frontend interfaces.
Integrate backend APIs.
Create reusable components.
Collaborate with backend developers and product teams.

Experience:
0-2 years accepted for internship or junior developer role.
"""
)

write_txt(
    BASE_DIR / "employee_queries" / "leave_query_sample.txt",
    """
Employee Query:
I am currently in probation. Can I take paid leave next week for a family function?
Please confirm whether probation employees are eligible for paid leave.

Employee Name: Rahul Shah
Department: Sales
Date: 2025-05-18
"""
)

write_txt(
    BASE_DIR / "disputes" / "salary_deduction_dispute.txt",
    """
Dispute Case:
Employee claims salary deduction was unfair because leave was already informed to manager.

Employee Name: Meera Joshi
Issue Type: Attendance Deduction
Employee Claim:
I had informed my manager about my absence on May 12, 2025, but my salary was deducted.

HR Notes:
Attendance system shows absent.
No approved leave application found in HR portal.
Manager confirmation pending.

Required Action:
Review attendance policy and draft HR response.
"""
)

# -----------------------------
# DOCX FILES
# -----------------------------

write_docx(
    BASE_DIR / "resumes" / "priya_nair_resume.docx",
    "Priya Nair Resume",
    [
        "Email: priya.nair@email.com",
        "Phone: +91 9123456780",
        "Role: Frontend Developer",
        "Summary: Frontend developer skilled in React, Tailwind CSS, responsive design, and API integration.",
        "Skills: React, JavaScript, Tailwind CSS, HTML, CSS, REST APIs, Git, Figma.",
        "Projects: Candidate Tracker UI, HRFlow mock dashboard, employee portal redesign.",
        "Experience: Frontend Intern at PixelStack Studio. Built reusable components and integrated backend APIs.",
        "Education: B.E. Computer Engineering.",
    ],
)

write_docx(
    BASE_DIR / "onboarding" / "joining_checklist.docx",
    "Employee Joining Checklist",
    [
        "Documents required before joining:",
        "1. Resume",
        "2. Government ID proof",
        "3. Address proof",
        "4. PAN card",
        "5. Bank account details",
        "6. Education certificates",
        "7. Previous experience letter, if applicable",
        "8. Signed offer letter",
        "HR must verify all documents before onboarding completion.",
    ],
)

# -----------------------------
# PDF FILES
# -----------------------------

write_pdf(
    BASE_DIR / "hr_policies" / "leave_policy.pdf",
    "HR Leave Policy",
    [
        "Annual paid leave entitlement: 18 days per calendar year.",
        "Employees under probation are eligible for 1 paid leave per month after 30 days of joining.",
        "Uninformed absence for more than 2 consecutive days may require HR review.",
        "Leave requests must be submitted at least 3 working days in advance unless emergency.",
        "Manager approval is required for all planned leaves.",
        "Leave without approval may be treated as unpaid leave.",
    ],
)

write_pdf(
    BASE_DIR / "hr_policies" / "attendance_policy.pdf",
    "Attendance Policy",
    [
        "Office timing is 10:00 AM to 7:00 PM.",
        "Employees arriving after 10:15 AM will receive a late mark.",
        "More than 3 late marks in a month require HR follow-up.",
        "Missing punch-in or punch-out must be regularized within 48 hours.",
        "Uninformed absence may lead to salary deduction.",
        "Repeated attendance violations may lead to warning letter.",
    ],
)

write_pdf(
    BASE_DIR / "job_descriptions" / "fullstack_developer_jd.pdf",
    "Full Stack Developer Job Description",
    [
        "Required skills: React, Node.js, Express.js, MongoDB, REST APIs, Git.",
        "Preferred skills: Next.js, FastAPI, authentication, deployment, Docker.",
        "Responsibilities include building frontend screens and backend APIs.",
        "Candidate should understand database design and API integration.",
        "Good project explanation and problem-solving ability are required.",
    ],
)

# -----------------------------
# CSV FILE
# -----------------------------

attendance_headers = [
    "employee_id",
    "employee_name",
    "department",
    "date",
    "check_in",
    "check_out",
    "status",
    "late_minutes",
]

attendance_rows = [
    ["E001", "Rahul Shah", "Sales", "2025-05-12", "10:22", "19:05", "Present", "7"],
    ["E002", "Meera Joshi", "Marketing", "2025-05-12", "", "", "Absent", ""],
    ["E003", "Ankit Rao", "Engineering", "2025-05-12", "10:05", "19:10", "Present", "0"],
    ["E004", "Sneha Iyer", "HR", "2025-05-12", "10:35", "19:00", "Present", "20"],
    ["E001", "Rahul Shah", "Sales", "2025-05-13", "10:28", "19:02", "Present", "13"],
    ["E004", "Sneha Iyer", "HR", "2025-05-13", "10:41", "19:12", "Present", "26"],
    ["E002", "Meera Joshi", "Marketing", "2025-05-13", "10:11", "19:00", "Present", "0"],
]

write_csv(
    BASE_DIR / "attendance" / "may_attendance_sample.csv",
    attendance_headers,
    attendance_rows,
)

# -----------------------------
# XLSX FILE
# -----------------------------

write_xlsx(
    BASE_DIR / "attendance" / "may_attendance_sample.xlsx",
    attendance_headers,
    attendance_rows,
)

# -----------------------------
# EXTRA GENERAL TXT
# -----------------------------

write_txt(
    BASE_DIR / "general" / "company_code_of_conduct.txt",
    """
Company Code of Conduct

Employees must maintain professional behavior, respect colleagues, protect company data,
avoid harassment, follow attendance policies, and report conflicts of interest.

Violation of company conduct may result in HR review, warning letter, or disciplinary action.
"""
)

print("")
print("Dummy HRFlow test files created successfully.")
print("")
print("Created formats:")
print("- TXT")
print("- DOCX")
print("- PDF")
print("- CSV")
print("- XLSX")
print("")
print(f"Location: {BASE_DIR}")
